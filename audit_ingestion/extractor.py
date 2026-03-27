"""
audit_ingestion_v03/audit_ingestion/extractor.py
Unified extraction engine.

Priority order for PDFs:
  1. pdfplumber (best for tables, digital PDFs)
  2. PyPDF2 (pure Python fallback, consistent cross-platform)
  3. extractous + Tesseract OCR (complex/scanned PDFs)
  4. PyMuPDF + pytesseract OCR (last resort before vision)
  5. OpenAI vision (image-based PDFs with no text layer)

For non-PDFs:
  - CSV/Excel/TXT: direct pandas/text read
  - DOCX: python-docx
"""
from __future__ import annotations
import logging
import io
from pathlib import Path
from typing import Optional
from dataclasses import dataclass, field

logger = logging.getLogger(__name__)

# Minimum chars per page to consider extraction acceptable
MIN_CHARS_PER_PAGE = 150
MIN_TOTAL_CHARS    = 300


@dataclass
class RawExtract:
    """Raw output from any extractor."""
    text:        Optional[str]
    tables:      list[dict]      = field(default_factory=list)
    page_count:  int             = 0
    extractor:   str             = "none"
    confidence:  float           = 0.0
    warnings:    list[str]       = field(default_factory=list)
    errors:      list[str]       = field(default_factory=list)

    @property
    def chars_per_page(self) -> float:
        if not self.page_count:
            return 0.0
        return len(self.text or "") / self.page_count

    @property
    def is_sufficient(self) -> bool:
        return (
            bool(self.text)
            and len(self.text) >= MIN_TOTAL_CHARS
            and self.chars_per_page >= MIN_CHARS_PER_PAGE
        )


# ── PDF Extractors ────────────────────────────────────────────────────────────

def _try_pdfplumber(path: str) -> RawExtract:
    try:
        import pdfplumber
    except ImportError:
        return RawExtract(text=None, extractor="pdfplumber", errors=["pdfplumber not installed"])

    warnings = []
    all_text = []
    tables = []

    try:
        with pdfplumber.open(path) as pdf:
            page_count = len(pdf.pages)
            for page_num, page in enumerate(pdf.pages[:200]):
                text = page.extract_text()
                if text:
                    all_text.append(text)

                page_tables = page.extract_tables() or []
                for ti, tbl in enumerate(page_tables):
                    if not tbl or len(tbl) < 2:
                        continue
                    headers = tbl[0]
                    valid_headers = sum(1 for h in headers if h and str(h).strip())
                    if valid_headers >= 2:
                        clean_headers = [str(h).strip() if h else f"Col_{j}" for j, h in enumerate(headers)]
                        rows = [{clean_headers[ci] if ci < len(clean_headers) else f"Col_{ci}": cell
                                 for ci, cell in enumerate(row)} for row in tbl[1:]]
                        tables.append({"page": page_num + 1, "headers": clean_headers,
                                       "rows": rows, "row_count": len(rows)})

        raw_text = "\n\n".join(all_text) or None
        confidence = 1.0 if raw_text and tables else (0.7 if raw_text else 0.0)

        return RawExtract(
            text=raw_text, tables=tables, page_count=page_count,
            extractor="pdfplumber", confidence=confidence, warnings=warnings
        )
    except Exception as e:
        return RawExtract(text=None, extractor="pdfplumber", errors=[str(e)])


def _try_pypdf2(path: str) -> RawExtract:
    try:
        from PyPDF2 import PdfReader
    except ImportError:
        return RawExtract(text=None, extractor="pypdf2", errors=["PyPDF2 not installed"])

    try:
        with open(path, "rb") as f:
            pdf_bytes = f.read()
        reader = PdfReader(io.BytesIO(pdf_bytes))
        page_count = len(reader.pages)
        texts = []
        for page in reader.pages:
            t = page.extract_text() or ""
            if t.strip():
                texts.append(t)
        raw_text = "\n\n".join(texts) or None
        confidence = 0.75 if raw_text else 0.0
        return RawExtract(text=raw_text, page_count=page_count,
                          extractor="pypdf2", confidence=confidence)
    except Exception as e:
        return RawExtract(text=None, extractor="pypdf2", errors=[str(e)])


def _try_extractous(path: str) -> RawExtract:
    try:
        from extractous import Extractor, TesseractOcrConfig
        extractor = (
            Extractor()
            .set_extract_string_max_length(2_000_000)
            .set_ocr_config(TesseractOcrConfig().set_language("eng"))
        )
        text, metadata = extractor.extract_file_to_string(str(path))
        meta_dict = dict(metadata) if metadata else {}

        # Normalize mime type
        raw_mime = meta_dict.get("Content-Type", "")
        if isinstance(raw_mime, list):
            raw_mime = raw_mime[0] if raw_mime else ""

        page_count = None
        for key in ["xmpTPg:NPages", "meta:page-count", "Page-Count"]:
            if key in meta_dict:
                try:
                    page_count = int(meta_dict[key])
                    break
                except (ValueError, TypeError):
                    pass

        raw_text = text.strip() if text and text.strip() else None
        confidence = 0.85 if raw_text and len(raw_text) > 500 else (0.4 if raw_text else 0.0)

        return RawExtract(
            text=raw_text, page_count=page_count or 0,
            extractor="extractous", confidence=confidence
        )
    except ImportError:
        return RawExtract(text=None, extractor="extractous", errors=["extractous not installed"])
    except Exception as e:
        return RawExtract(text=None, extractor="extractous", errors=[str(e)])


def _try_ocr(path: str) -> RawExtract:
    """PyMuPDF + pytesseract OCR fallback."""
    try:
        import fitz
        import pytesseract
        from PIL import Image
    except ImportError:
        return RawExtract(text=None, extractor="ocr", errors=["PyMuPDF/pytesseract not installed"])

    try:
        doc = fitz.open(path)
        page_count = min(len(doc), 200)
        page_texts = []

        for i in range(page_count):
            page = doc[i]
            pix = page.get_pixmap(dpi=250, alpha=False)
            mode = "RGB" if pix.n < 4 else "RGBA"
            img = Image.frombytes(mode, [pix.width, pix.height], pix.samples)
            text = pytesseract.image_to_string(img, lang="eng", config="--oem 1 --psm 6")
            if text and text.strip():
                page_texts.append(text)

        raw_text = "\n\n".join(page_texts).strip() or None
        confidence = 0.85 if raw_text and len(raw_text) >= 2000 else (0.55 if raw_text else 0.0)

        return RawExtract(
            text=raw_text, page_count=page_count,
            extractor="ocr", confidence=confidence
        )
    except Exception as e:
        return RawExtract(text=None, extractor="ocr", errors=[str(e)])


def _try_vision(path: str, provider) -> RawExtract:
    """OpenAI GPT-4o vision fallback — reads PDF pages as images."""
    if not hasattr(provider, "extract_text_from_pdf_vision"):
        return RawExtract(text=None, extractor="vision", errors=["Provider does not support vision"])

    try:
        with open(path, "rb") as f:
            pdf_bytes = f.read()
        text = provider.extract_text_from_pdf_vision(pdf_bytes, max_pages=6)
        raw_text = text.strip() if text and text.strip() else None
        confidence = 0.90 if raw_text and len(raw_text) > 500 else (0.5 if raw_text else 0.0)
        return RawExtract(
            text=raw_text, page_count=0,
            extractor="vision", confidence=confidence
        )
    except Exception as e:
        return RawExtract(text=None, extractor="vision", errors=[str(e)])


# ── Non-PDF Extractors ────────────────────────────────────────────────────────

def _extract_direct(path: str) -> RawExtract:
    """CSV, Excel, TXT direct extraction."""
    p = Path(path)
    ext = p.suffix.lower()

    try:
        if ext in (".csv", ".tsv"):
            import pandas as pd
            df = pd.read_csv(path)
            text = df.to_string(index=False)
            return RawExtract(text=text, extractor="direct", confidence=1.0,
                              tables=[{"headers": df.columns.tolist(),
                                       "rows": df.head(100).to_dict("records"),
                                       "row_count": len(df)}])

        elif ext in (".xlsx", ".xls"):
            import pandas as pd
            df = pd.read_excel(path)
            text = df.to_string(index=False)
            return RawExtract(text=text, extractor="direct", confidence=1.0,
                              tables=[{"headers": df.columns.tolist(),
                                       "rows": df.head(100).to_dict("records"),
                                       "row_count": len(df)}])

        elif ext == ".txt":
            text = p.read_text(encoding="utf-8", errors="replace")
            return RawExtract(text=text, extractor="direct", confidence=1.0)

        elif ext == ".docx":
            from docx import Document
            doc = Document(path)
            text = "\n".join(p.text for p in doc.paragraphs if p.text.strip())
            return RawExtract(text=text, extractor="direct", confidence=0.9)

        else:
            return RawExtract(text=None, extractor="direct",
                              errors=[f"Unsupported file type: {ext}"])

    except Exception as e:
        return RawExtract(text=None, extractor="direct", errors=[str(e)])


# ── Main Extract Function ─────────────────────────────────────────────────────

def extract(path: str, provider=None) -> RawExtract:
    """
    Extract text and tables from any supported file.
    For PDFs: tries multiple extractors in order, picks best result.
    Stops early if extraction is sufficient.
    """
    p = Path(path)
    ext = p.suffix.lower()

    # Non-PDF files go straight to direct extraction
    if ext not in (".pdf",):
        return _extract_direct(path)

    # PDF — try extractors in order, pick best
    candidates: list[RawExtract] = []

    # 1. pdfplumber — best for tables and digital PDFs
    result = _try_pdfplumber(path)
    candidates.append(result)
    if result.is_sufficient and result.tables:
        logger.info(f"pdfplumber sufficient ({len(result.text or '')} chars, {len(result.tables)} tables)")
        return result

    # 2. PyPDF2 — pure Python, consistent cross-platform
    result = _try_pypdf2(path)
    candidates.append(result)

    # Pick best so far
    best = max(candidates, key=lambda r: len(r.text or ""))
    if best.is_sufficient:
        logger.info(f"PyPDF2 sufficient ({len(best.text or '')} chars)")
        return best

    # 3. extractous + Tesseract OCR
    result = _try_extractous(path)
    candidates.append(result)
    best = max(candidates, key=lambda r: len(r.text or ""))
    if best.is_sufficient:
        logger.info(f"extractous sufficient ({len(best.text or '')} chars)")
        return best

    # 4. PyMuPDF + pytesseract OCR
    result = _try_ocr(path)
    candidates.append(result)
    best = max(candidates, key=lambda r: len(r.text or ""))
    if best.is_sufficient:
        logger.info(f"OCR sufficient ({len(best.text or '')} chars)")
        return best

    # 5. Vision — last resort, requires provider
    if provider is not None:
        result = _try_vision(path, provider)
        candidates.append(result)
        best = max(candidates, key=lambda r: len(r.text or ""))
        logger.info(f"Vision used ({len(best.text or '')} chars)")
        return best

    # Return whatever we have even if insufficient
    best = max(candidates, key=lambda r: len(r.text or ""))
    best.warnings.append("Extraction quality below threshold — review recommended")
    return best
